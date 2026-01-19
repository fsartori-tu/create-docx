from flask import Flask, request, send_file
from io import BytesIO
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
from docx.enum.text import WD_BREAK
import os

app = Flask(__name__)
PLACEHOLDER = "{{COMPLETAR}}"

# -------------------------
# Helpers
# -------------------------

def format_name_from_email(value: str) -> str:
    if not value:
        return ""
    value = value.strip()
    if "@" not in value:
        return value
    local = value.split("@")[0]
    parts = [p.strip() for p in local.split(".") if p.strip()]
    return " ".join(p.capitalize() for p in parts)

def safe_text(value: str) -> str:
    v = (value or "").strip()
    return v if v else PLACEHOLDER

def insert_paragraph_after(paragraph: Paragraph, text: str = "", style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para

def remove_paragraph(paragraph: Paragraph) -> None:
    p = paragraph._p
    p.getparent().remove(p)

def find_paragraph_index(doc: Document, target_text: str) -> int:
    t = (target_text or "").strip().lower()
    for i, p in enumerate(doc.paragraphs):
        if (p.text or "").strip().lower() == t:
            return i
    return -1

def is_heading1(p: Paragraph) -> bool:
    """Detecta Heading 1 o Título 1 de forma robusta"""
    try:
        name = (p.style.name or "").lower()
        return "heading 1" in name or "título 1" in name or "titulo 1" in name
    except:
        return False

def set_labeled_line(doc: Document, label: str, value: str) -> bool:
    label_norm = (label or "").strip()
    for p in doc.paragraphs:
        txt = (p.text or "").strip()
        if txt.startswith(label_norm):
            p.text = f"{label_norm} {value if value else PLACEHOLDER}"
            return True
    return False

def set_run_font_size(paragraph: Paragraph, size_pt: int) -> None:
    for r in paragraph.runs:
        r.font.size = Pt(size_pt)

def add_page_break_after(paragraph: Paragraph) -> Paragraph:
    p_break = insert_paragraph_after(paragraph, "")
    run = p_break.add_run()
    run.add_break(WD_BREAK.PAGE)
    return p_break

def clear_section_body(doc: Document, start_idx: int) -> int:
    """
    Limpia los párrafos entre el título (start_idx) y el próximo Heading 1.
    Devuelve el end_idx (inicio del próximo Heading1 o fin).
    """
    end_idx = -1
    for j in range(start_idx + 1, len(doc.paragraphs)):
        if is_heading1(doc.paragraphs[j]):
            end_idx = j
            break
    if end_idx == -1:
        end_idx = len(doc.paragraphs)

    for p in reversed(doc.paragraphs[start_idx + 1:end_idx]):
        remove_paragraph(p)

    return end_idx

def insert_multiline_text(cursor: Paragraph, text: str, style: str = "Normal", font_size: int = 11) -> Paragraph:
    """
    Inserta texto respetando saltos de línea: cada línea => un párrafo.
    Devuelve el último párrafo insertado.
    """
    lines = (text or "").splitlines()
    lines = [ln.rstrip() for ln in lines]

    if not any(ln.strip() for ln in lines):
        p = insert_paragraph_after(cursor, PLACEHOLDER, style=style)
        set_run_font_size(p, font_size)
        return p

    last = cursor
    for ln in lines:
        p = insert_paragraph_after(last, ln if ln.strip() else "", style=style)
        set_run_font_size(p, font_size)
        last = p

    return last

# -------------------------
# Secciones
# -------------------------

def rebuild_summary_section(doc: Document, resumen: str | None) -> None:
    """
    Busca la sección 'Resúmen' (Heading 1) y reemplaza su contenido.
    """
    start_idx = find_paragraph_index(doc, "Resúmen")
    if start_idx == -1:
        return

    clear_section_body(doc, start_idx)

    cursor = doc.paragraphs[start_idx]
    insert_multiline_text(cursor, (resumen or "").strip(), style="Normal", font_size=11)

def rebuild_self_feedback_section(doc: Document, auto: dict | None) -> None:
    start_idx = find_paragraph_index(doc, "Autoevaluación")
    if start_idx == -1:
        return

    clear_section_body(doc, start_idx)

    cursor = doc.paragraphs[start_idx]
    has_content = auto and any(auto.get(k) for k in ["positivos", "mejorar", "algo_mas"])

    if not has_content:
        p = insert_paragraph_after(cursor, "No se registró autoevaluación.", style="Normal")
        set_run_font_size(p, 11)
        cursor = p
    else:
        mapping = [
            ("1. Aspectos positivos", auto.get("positivos")),
            ("2. Aspectos a mejorar", auto.get("mejorar")),
            ("3. Algo más que quieras compartir.", auto.get("algo_mas")),
        ]
        for label, content in mapping:
            p_lab = insert_paragraph_after(cursor, "", style="Normal")
            run = p_lab.add_run(label)
            run.bold = True
            cursor = p_lab

            p_txt = insert_paragraph_after(cursor, safe_text(content), style="Normal")
            set_run_font_size(p_txt, 11)
            cursor = p_txt

    add_page_break_after(cursor)

def rebuild_feedback_section(doc: Document, evaluaciones: list) -> None:
    start_idx = find_paragraph_index(doc, "Feedback recibido")
    if start_idx == -1:
        return

    clear_section_body(doc, start_idx)

    cursor = doc.paragraphs[start_idx]
    if not evaluaciones:
        p = insert_paragraph_after(cursor, PLACEHOLDER, style="Normal")
        set_run_font_size(p, 11)
        return

    for idx, ev in enumerate(evaluaciones):
        evaluador = format_name_from_email(ev.get("evaluador", "")) or PLACEHOLDER
        p_eval = insert_paragraph_after(cursor, evaluador, style="Heading 2")
        cursor = p_eval

        items = [
            ("1. Aspectos positivos", ev.get("positivos")),
            ("2. Aspectos a mejorar", ev.get("mejorar")),
            ("3. Algo más que quieras compartir.", ev.get("algo_mas")),
        ]

        for q, a in items:
            p_q = insert_paragraph_after(cursor, "", style="Normal")
            p_q.add_run(q).bold = True
            cursor = p_q

            p_a = insert_paragraph_after(cursor, safe_text(a), style="Normal")
            set_run_font_size(p_a, 11)
            cursor = p_a

        if idx < len(evaluaciones) - 1:
            cursor = add_page_break_after(cursor)

# -------------------------
# Ruta Principal
# -------------------------

@app.post("/generate")
def generate():
    try:
        data = request.json or {}

        raw_evaluado = (data.get("evaluado") or "").strip()
        evaluado = format_name_from_email(raw_evaluado) or PLACEHOLDER

        mes_ano = (data.get("mes_ano") or "").strip() or PLACEHOLDER

        # Soporta ambos: "autoev" (tu output limpio) o "autoevaluacion" (legacy)
        auto = data.get("autoev") or data.get("autoevaluacion")
        evs = data.get("evaluaciones") or []

        # NUEVO: resumen
        resumen = (data.get("resumen") or "").strip()

        template_path = os.environ.get("TEMPLATE_PATH", "template.docx")
        if not os.path.exists(template_path):
            return {"error": f"Template no encontrado en {template_path}"}, 404

        doc = Document(template_path)

        set_labeled_line(doc, "Nombre:", evaluado)
        set_labeled_line(doc, "Periodo evaluado:", mes_ano)

        # Inserta resumen en sección "Resúmen"
        rebuild_summary_section(doc, resumen)

        rebuild_self_feedback_section(doc, auto)
        rebuild_feedback_section(doc, evs)

        bio = BytesIO()
        doc.save(bio)
        bio.seek(0)

        filename = f"Performance Review – {evaluado} – {mes_ano}.docx"

        return send_file(
            bio,
            as_attachment=True,
            download_name=filename,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    except Exception as e:
        print(f"Error generando DOCX: {str(e)}")
        return {"error": str(e)}, 500

if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port)
